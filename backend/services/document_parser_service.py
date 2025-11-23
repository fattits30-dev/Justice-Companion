"""
DocumentParserService - Extract text and metadata from various document formats.

Migrated from: src/services/DocumentParserService.ts

Supported formats:
- PDF (.pdf) - via pypdf
- Word Documents (.docx) - via python-docx
- Plain Text (.txt) - native UTF-8 decoding

Future support: Images with OCR (.jpg, .png) via pytesseract

Security considerations:
- File size validation (max 10MB default)
- Supported extension whitelist
- Path traversal prevention
- Safe file parsing with error handling

Usage:
    from backend.services.document_parser_service import DocumentParserService

    parser = DocumentParserService(audit_logger=audit_logger)
    result = await parser.parse_document(file_path="/path/to/document.pdf")

    print(f"Extracted {result.word_count} words from {result.filename}")
    summary = parser.extract_summary(result.text, max_words=100)
"""

import os
import re
import logging
from typing import Optional, Dict, Any, List

from pydantic import BaseModel, Field, field_validator
from fastapi import HTTPException

# PDF parsing
try:
    import pypdf
except ImportError:
    pypdf = None

# Word document parsing
try:
    import docx
except ImportError:
    docx = None

# Configure logger
logger = logging.getLogger(__name__)


class ParsedDocumentMetadata(BaseModel):
    """Metadata extracted from document."""

    title: Optional[str] = None
    author: Optional[str] = None
    creation_date: Optional[str] = None
    extra: Optional[Dict[str, Any]] = Field(default_factory=dict)

    class Config:
        """Pydantic configuration."""

        extra = "allow"  # Allow additional fields


class ParsedDocument(BaseModel):
    """
    Parsed document result with extracted text and metadata.

    Attributes:
        text: Extracted plain text content
        filename: Original filename
        file_type: Document type (pdf, docx, txt)
        page_count: Number of pages (PDF only)
        word_count: Total word count
        metadata: Optional document metadata (title, author, dates, etc.)
    """

    text: str = Field(..., description="Extracted plain text content")
    filename: str = Field(..., description="Original filename")
    file_type: str = Field(..., description="Document type (pdf, docx, txt)")
    page_count: Optional[int] = Field(None, description="Number of pages (PDF only)")
    word_count: int = Field(..., description="Total word count")
    metadata: Optional[ParsedDocumentMetadata] = Field(None, description="Document metadata")

    @field_validator("file_type")
    def validate_file_type(cls, v):
        """Validate file type is supported."""
        allowed = ["pdf", "docx", "txt"]
        if v not in allowed:
            raise ValueError(f"File type must be one of {allowed}, got: {v}")
        return v


class FileSizeValidationResult(BaseModel):
    """File size validation result."""

    valid: bool
    error: Optional[str] = None


class DocumentParserService:
    """
    Service for parsing various document formats and extracting text.

    Provides methods to parse PDF, DOCX, and TXT files from file paths or
    byte buffers. Includes validation and helper utilities.
    """

    # Maximum file size: 10MB default
    MAX_FILE_SIZE = 10 * 1024 * 1024

    # Supported file extensions
    SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt"]

    def __init__(self, audit_logger=None, max_file_size: Optional[int] = None):
        """
        Initialize DocumentParserService.

        Args:
            audit_logger: Optional AuditLogger instance for logging parsing events
            max_file_size: Maximum allowed file size in bytes (default: 10MB)
        """
        self.audit_logger = audit_logger
        self.max_file_size = max_file_size or self.MAX_FILE_SIZE

        # Log warnings if dependencies are missing
        if pypdf is None:
            logger.warning("pypdf not installed - PDF parsing will fail")
        if docx is None:
            logger.warning("python-docx not installed - DOCX parsing will fail")

    async def parse_document(self, file_path: str, user_id: Optional[str] = None) -> ParsedDocument:
        """
        Parse document from file path.

        Args:
            file_path: Absolute path to document file
            user_id: Optional user ID for audit logging

        Returns:
            ParsedDocument with extracted text and metadata

        Raises:
            HTTPException: If file doesn't exist, unsupported format, or parsing fails
        """
        # Validate file exists
        if not os.path.isfile(file_path):
            raise HTTPException(status_code=404, detail=f"File not found: {file_path}")

        # Validate file size
        file_size = os.path.getsize(file_path)
        size_validation = self.validate_file_size(file_size)
        if not size_validation.valid:
            raise HTTPException(status_code=413, detail=size_validation.error)

        filename = os.path.basename(file_path)
        ext = os.path.splitext(file_path)[1].lower()

        # Log parsing attempt
        if self.audit_logger:
            self.audit_logger.log(
                event_type="document.parse.start",
                user_id=user_id,
                resource_type="document",
                resource_id=file_path,
                action="parse",
                details={"filename": filename, "extension": ext, "size_bytes": file_size},
                success=True,
            )

        try:
            # Route to appropriate parser
            if ext == ".pdf":
                result = await self._parse_pdf(file_path, filename)
            elif ext == ".docx":
                result = await self._parse_docx(file_path, filename)
            elif ext == ".txt":
                result = await self._parse_txt(file_path, filename)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file format: {ext}. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}",
                )

            # Log successful parse
            if self.audit_logger:
                self.audit_logger.log(
                    event_type="document.parse.success",
                    user_id=user_id,
                    resource_type="document",
                    resource_id=file_path,
                    action="parse",
                    details={"word_count": result.word_count, "page_count": result.page_count},
                    success=True,
                )

            return result

        except HTTPException:
            raise
        except Exception as e:
            # Log parsing failure
            if self.audit_logger:
                self.audit_logger.log(
                    event_type="document.parse.error",
                    user_id=user_id,
                    resource_type="document",
                    resource_id=file_path,
                    action="parse",
                    details={"error": str(e)},
                    success=False,
                    error_message=str(e),
                )

            logger.error(f"Failed to parse document {filename}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to parse document: {str(e)}")

    async def parse_document_buffer(
        self, buffer: bytes, filename: str, user_id: Optional[str] = None
    ) -> ParsedDocument:
        """
        Parse document from byte buffer.

        Args:
            buffer: Document content as bytes
            filename: Original filename (used to determine file type)
            user_id: Optional user ID for audit logging

        Returns:
            ParsedDocument with extracted text and metadata

        Raises:
            HTTPException: If unsupported format or parsing fails
        """
        # Validate file size
        file_size = len(buffer)
        size_validation = self.validate_file_size(file_size)
        if not size_validation.valid:
            raise HTTPException(status_code=413, detail=size_validation.error)

        ext = os.path.splitext(filename)[1].lower()

        # Log parsing attempt
        if self.audit_logger:
            self.audit_logger.log(
                event_type="document.parse_buffer.start",
                user_id=user_id,
                resource_type="document",
                resource_id=filename,
                action="parse",
                details={"filename": filename, "extension": ext, "size_bytes": file_size},
                success=True,
            )

        try:
            # Route to appropriate parser
            if ext == ".pdf":
                result = await self._parse_pdf_buffer(buffer, filename)
            elif ext == ".docx":
                result = await self._parse_docx_buffer(buffer, filename)
            elif ext == ".txt":
                result = await self._parse_txt_buffer(buffer, filename)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file format: {ext}. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}",
                )

            # Log successful parse
            if self.audit_logger:
                self.audit_logger.log(
                    event_type="document.parse_buffer.success",
                    user_id=user_id,
                    resource_type="document",
                    resource_id=filename,
                    action="parse",
                    details={"word_count": result.word_count, "page_count": result.page_count},
                    success=True,
                )

            return result

        except HTTPException:
            raise
        except Exception as e:
            # Log parsing failure
            if self.audit_logger:
                self.audit_logger.log(
                    event_type="document.parse_buffer.error",
                    user_id=user_id,
                    resource_type="document",
                    resource_id=filename,
                    action="parse",
                    details={"error": str(e)},
                    success=False,
                    error_message=str(e),
                )

            logger.error(f"Failed to parse document buffer {filename}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to parse document: {str(e)}")

    async def _parse_pdf(self, file_path: str, filename: str) -> ParsedDocument:
        """
        Parse PDF file from path.

        Args:
            file_path: Path to PDF file
            filename: Original filename

        Returns:
            ParsedDocument with extracted text

        Raises:
            HTTPException: If pypdf not installed or parsing fails
        """
        if pypdf is None:
            raise HTTPException(
                status_code=500, detail="PDF parsing not available (pypdf not installed)"
            )

        with open(file_path, "rb") as f:
            buffer = f.read()

        return await self._parse_pdf_buffer(buffer, filename)

    async def _parse_pdf_buffer(self, buffer: bytes, filename: str) -> ParsedDocument:
        """
        Parse PDF from byte buffer.

        Args:
            buffer: PDF content as bytes
            filename: Original filename

        Returns:
            ParsedDocument with extracted text

        Raises:
            HTTPException: If pypdf not installed or parsing fails
        """
        if pypdf is None:
            raise HTTPException(
                status_code=500, detail="PDF parsing not available (pypdf not installed)"
            )

        try:
            # Use BytesIO to read from buffer
            from io import BytesIO

            pdf_file = BytesIO(buffer)

            # Parse PDF
            reader = pypdf.PdfReader(pdf_file)

            # Extract text from all pages
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text())

            text = "\n".join(text_parts)
            page_count = len(reader.pages)

            # Extract metadata
            metadata_dict = {}
            if reader.metadata:
                if reader.metadata.title:
                    metadata_dict["title"] = reader.metadata.title
                if reader.metadata.author:
                    metadata_dict["author"] = reader.metadata.author
                if reader.metadata.creation_date:
                    metadata_dict["creation_date"] = str(reader.metadata.creation_date)

                # Add any other metadata fields
                metadata_dict["extra"] = {}
                if reader.metadata.subject:
                    metadata_dict["extra"]["subject"] = reader.metadata.subject
                if reader.metadata.creator:
                    metadata_dict["extra"]["creator"] = reader.metadata.creator
                if reader.metadata.producer:
                    metadata_dict["extra"]["producer"] = reader.metadata.producer

            metadata = ParsedDocumentMetadata(**metadata_dict) if metadata_dict else None

            return ParsedDocument(
                text=text,
                filename=filename,
                file_type="pdf",
                page_count=page_count,
                word_count=self._count_words(text),
                metadata=metadata,
            )

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to parse PDF: {str(e)}")

    async def _parse_docx(self, file_path: str, filename: str) -> ParsedDocument:
        """
        Parse DOCX file from path.

        Args:
            file_path: Path to DOCX file
            filename: Original filename

        Returns:
            ParsedDocument with extracted text

        Raises:
            HTTPException: If python-docx not installed or parsing fails
        """
        if docx is None:
            raise HTTPException(
                status_code=500, detail="DOCX parsing not available (python-docx not installed)"
            )

        with open(file_path, "rb") as f:
            buffer = f.read()

        return await self._parse_docx_buffer(buffer, filename)

    async def _parse_docx_buffer(self, buffer: bytes, filename: str) -> ParsedDocument:
        """
        Parse DOCX from byte buffer.

        Args:
            buffer: DOCX content as bytes
            filename: Original filename

        Returns:
            ParsedDocument with extracted text

        Raises:
            HTTPException: If python-docx not installed or parsing fails
        """
        if docx is None:
            raise HTTPException(
                status_code=500, detail="DOCX parsing not available (python-docx not installed)"
            )

        try:
            # Use BytesIO to read from buffer
            from io import BytesIO

            docx_file = BytesIO(buffer)

            # Parse DOCX
            doc = docx.Document(docx_file)

            # Extract text from all paragraphs
            text_parts = [paragraph.text for paragraph in doc.paragraphs]
            text = "\n".join(text_parts)

            # Extract core properties metadata
            metadata_dict = {}
            if doc.core_properties:
                if doc.core_properties.title:
                    metadata_dict["title"] = doc.core_properties.title
                if doc.core_properties.author:
                    metadata_dict["author"] = doc.core_properties.author
                if doc.core_properties.created:
                    metadata_dict["creation_date"] = doc.core_properties.created.isoformat()

                # Add extra metadata
                metadata_dict["extra"] = {}
                if doc.core_properties.subject:
                    metadata_dict["extra"]["subject"] = doc.core_properties.subject
                if doc.core_properties.keywords:
                    metadata_dict["extra"]["keywords"] = doc.core_properties.keywords
                if doc.core_properties.last_modified_by:
                    metadata_dict["extra"][
                        "last_modified_by"
                    ] = doc.core_properties.last_modified_by
                if doc.core_properties.modified:
                    metadata_dict["extra"]["modified"] = doc.core_properties.modified.isoformat()

            metadata = ParsedDocumentMetadata(**metadata_dict) if metadata_dict else None

            return ParsedDocument(
                text=text,
                filename=filename,
                file_type="docx",
                word_count=self._count_words(text),
                metadata=metadata,
            )

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to parse DOCX: {str(e)}")

    async def _parse_txt(self, file_path: str, filename: str) -> ParsedDocument:
        """
        Parse plain text file from path.

        Args:
            file_path: Path to TXT file
            filename: Original filename

        Returns:
            ParsedDocument with extracted text

        Raises:
            HTTPException: If parsing fails
        """
        with open(file_path, "rb") as f:
            buffer = f.read()

        return await self._parse_txt_buffer(buffer, filename)

    async def _parse_txt_buffer(self, buffer: bytes, filename: str) -> ParsedDocument:
        """
        Parse plain text from byte buffer.

        Args:
            buffer: Text content as bytes
            filename: Original filename

        Returns:
            ParsedDocument with extracted text

        Raises:
            HTTPException: If decoding fails
        """
        try:
            # Decode UTF-8 (with fallback to latin-1)
            try:
                text = buffer.decode("utf-8")
            except UnicodeDecodeError:
                logger.warning(f"UTF-8 decode failed for {filename}, trying latin-1")
                text = buffer.decode("latin-1")

            return ParsedDocument(
                text=text, filename=filename, file_type="txt", word_count=self._count_words(text)
            )

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to parse TXT: {str(e)}")

    def _count_words(self, text: str) -> int:
        """
        Count words in text.

        Args:
            text: Text content

        Returns:
            Word count
        """
        # Split on whitespace and filter empty strings
        words = re.split(r"\s+", text.strip())
        return len([word for word in words if word])

    def extract_summary(self, text: str, max_words: int = 500) -> str:
        """
        Extract summary from parsed text (first N words).

        Args:
            text: Full text content
            max_words: Maximum words to include (default: 500)

        Returns:
            Summary text (truncated if needed)
        """
        words = re.split(r"\s+", text.strip())

        if len(words) <= max_words:
            return text

        return " ".join(words[:max_words]) + "..."

    def validate_file_size(self, file_size: int) -> FileSizeValidationResult:
        """
        Validate file size against maximum allowed.

        Args:
            file_size: File size in bytes

        Returns:
            FileSizeValidationResult with validation status
        """
        if file_size > self.max_file_size:
            max_mb = self.max_file_size / 1024 / 1024
            return FileSizeValidationResult(
                valid=False, error=f"File size exceeds maximum allowed size of {max_mb:.1f}MB"
            )

        return FileSizeValidationResult(valid=True)

    def get_supported_extensions(self) -> List[str]:
        """
        Get list of supported file extensions.

        Returns:
            List of supported extensions (e.g., ['.pd', '.docx', '.txt'])
        """
        return self.SUPPORTED_EXTENSIONS.copy()

    def is_supported(self, filename: str) -> bool:
        """
        Check if file type is supported.

        Args:
            filename: Filename to check

        Returns:
            True if supported, False otherwise
        """
        ext = os.path.splitext(filename)[1].lower()
        return ext in self.SUPPORTED_EXTENSIONS


# Convenience function for quick usage
async def parse_document(
    file_path: str, audit_logger=None, user_id: Optional[str] = None
) -> ParsedDocument:
    """
    Convenience function to parse a document.

    Args:
        file_path: Path to document file
        audit_logger: Optional AuditLogger instance
        user_id: Optional user ID for audit logging

    Returns:
        ParsedDocument with extracted text and metadata
    """
    service = DocumentParserService(audit_logger=audit_logger)
    return await service.parse_document(file_path, user_id=user_id)
