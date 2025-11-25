"""
DOCX Generator Service for Justice Companion

This module provides functionality for generating Microsoft Word documents (.docx)
from case data, evidence lists, timelines, and notes. It uses python-docx for
document generation with professional formatting, headers, footers, and page numbering.

Key Features:
- Generate comprehensive case summary reports
- Create evidence inventory reports
- Build timeline reports with event chronology
- Export case notes in structured format
- Professional document formatting with headers/footers
- Support for custom margins, spacing, and typography
- Automatic page breaks between sections
- Audit logging for all export operations

Author: Justice Companion Development Team
License: MIT
Python Version: 3.12+
"""

from datetime import datetime
from typing import Any, List, Optional
import io

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pydantic import BaseModel, Field, field_validator

# Type definitions matching TypeScript models

class TimelineEvent(BaseModel):
    """Timeline event representing a case milestone or deadline."""

    id: int
    case_id: int
    title: str
    description: Optional[str] = None
    event_date: str  # ISO 8601 date string
    event_type: str = Field(..., pattern="^(deadline|hearing|filing|milestone|other)$")
    completed: bool = False
    created_at: str
    updated_at: str

    @field_validator("event_date", "created_at", "updated_at")
    @classmethod
    def validate_iso_date(cls, v: str) -> str:
        """Validate ISO 8601 date format."""
        try:
            datetime.fromisoformat(v.replace("Z", "+00:00"))
            return v
        except ValueError:
            raise ValueError(f"Invalid ISO 8601 date format: {v}")

class Evidence(BaseModel):
    """Evidence item associated with a case."""

    id: int
    case_id: int
    title: str
    file_path: Optional[str] = None
    content: Optional[str] = None
    evidence_type: str = Field(..., pattern="^(document|photo|email|recording|note|witness)$")
    obtained_date: Optional[str] = None
    created_at: str
    updated_at: Optional[str] = None

class Note(BaseModel):
    """Case note or annotation."""

    id: int
    case_id: Optional[int] = None
    user_id: int
    title: Optional[str] = None
    content: str
    is_pinned: bool = False
    created_at: str
    updated_at: str

class Case(BaseModel):
    """Legal case record."""

    id: int
    title: str
    description: str
    status: str

class CaseExportData(BaseModel):
    """Complete case export data structure."""

    case: Case
    evidence: List[Evidence] = Field(default_factory=list)
    timeline: List[TimelineEvent] = Field(default_factory=list)
    notes: List[Note] = Field(default_factory=list)
    export_date: datetime
    exported_by: str

class EvidenceExportData(BaseModel):
    """Evidence inventory export data."""

    case_id: int
    case_title: str
    evidence: List[Evidence]
    export_date: datetime
    exported_by: str
    total_items: int

class TimelineExportData(BaseModel):
    """Timeline report export data."""

    case_id: int
    case_title: str
    events: List[TimelineEvent]
    export_date: datetime
    exported_by: str

class NotesExportData(BaseModel):
    """Case notes export data."""

    case_id: int
    case_title: str
    notes: List[Note]
    export_date: datetime
    exported_by: str
    total_notes: int

class DOCXGenerator:
    """
    DOCX document generator for legal case exports.

    This class provides methods to generate professionally formatted Word documents
    from case data, evidence inventories, timelines, and notes. All documents include
    headers, footers, page numbering, and consistent typography.

    Attributes:
        audit_logger: Optional audit logger for tracking export operations

    Example:
        ```python
        generator = DOCXGenerator(audit_logger=logger)

        # Generate case summary
        docx_bytes = await generator.generate_case_summary(case_data)

        # Save to file
        with open('case_summary.docx', 'wb') as f:
            f.write(docx_bytes)
        ```
    """

    # Constants for document formatting (matching TypeScript 1440 twips = 1 inch)
    MARGIN_TOP = Inches(1)
    MARGIN_RIGHT = Inches(1)
    MARGIN_BOTTOM = Inches(1)
    MARGIN_LEFT = Inches(1)

    FONT_SIZE_TITLE = Pt(24)
    FONT_SIZE_HEADING_1 = Pt(18)
    FONT_SIZE_HEADING_2 = Pt(14)
    FONT_SIZE_HEADING_3 = Pt(12)
    FONT_SIZE_BODY = Pt(11)
    FONT_SIZE_FOOTER = Pt(9)

    SPACING_AFTER_TITLE = Pt(10)
    SPACING_AFTER_HEADING = Pt(10)
    SPACING_AFTER_PARAGRAPH = Pt(5)

    def __init__(self, audit_logger: Optional[Any] = None):
        """
        Initialize DOCX generator.

        Args:
            audit_logger: Optional audit logger for tracking export operations
        """
        self.audit_logger = audit_logger

    async def generate_case_summary(self, case_data: CaseExportData) -> bytes:
        """
        Generate comprehensive case summary DOCX document.

        Creates a multi-section document including:
        - Case information (title, description, status)
        - Evidence inventory (if present)
        - Timeline of events (if present)
        - Case notes (if present)

        Each section is separated by page breaks and includes professional formatting.

        Args:
            case_data: Complete case export data structure

        Returns:
            bytes: DOCX document as bytes buffer

        Raises:
            Exception: If document generation fails

        Example:
            ```python
            case_data = CaseExportData(
                case=Case(id=1, title="Smith v. Jones", description="...", status="active"),
                evidence=[...],
                timeline=[...],
                notes=[...],
                export_date=datetime.now(),
                exported_by="user@example.com"
            )

            docx_bytes = await generator.generate_case_summary(case_data)
            ```
        """
        try:
            doc = Document()

            # Set document margins
            self._set_margins(doc)

            # Add header
            self._add_header(doc, f"Case: {case_data.case.title}")

            # Add footer with export info and page numbers
            self._add_footer(
                doc,
                f"Exported by {case_data.exported_by} on {case_data.export_date.strftime('%Y-%m-%d')}",
            )

            # Title section
            self._add_title(doc, "Case Summary")

            # Case information section
            self._add_case_info(doc, case_data)

            # Evidence section (if present)
            if case_data.evidence:
                doc.add_page_break()
                self._add_evidence_section(doc, case_data.evidence)

            # Timeline section (if present)
            if case_data.timeline:
                doc.add_page_break()
                self._add_timeline_section(doc, case_data.timeline)

            # Notes section (if present)
            if case_data.notes:
                doc.add_page_break()
                self._add_notes_section(doc, case_data.notes)

            # Log audit event
            if self.audit_logger:
                await self._log_audit(
                    "case_summary_generated", case_data.case.id, case_data.exported_by
                )

            # Convert to bytes
            return self._document_to_bytes(doc)

        except Exception as exc:
            if self.audit_logger:
                await self._log_audit(
                    "case_summary_failed", case_data.case.id, case_data.exported_by, error=str(e)
                )
            raise

    async def generate_evidence_list(self, evidence_data: EvidenceExportData) -> bytes:
        """
        Generate evidence inventory report DOCX document.

        Creates a detailed inventory of all evidence items including:
        - Case identification
        - Total evidence count
        - Individual evidence entries with metadata

        Args:
            evidence_data: Evidence export data structure

        Returns:
            bytes: DOCX document as bytes buffer

        Example:
            ```python
            evidence_data = EvidenceExportData(
                case_id=1,
                case_title="Smith v. Jones",
                evidence=[...],
                export_date=datetime.now(),
                exported_by="user@example.com",
                total_items=10
            )

            docx_bytes = await generator.generate_evidence_list(evidence_data)
            ```
        """
        try:
            doc = Document()
            self._set_margins(doc)

            # Add header and footer
            self._add_header(doc, f"Evidence Report: {evidence_data.case_title}")
            self._add_footer(
                doc,
                f"Exported by {evidence_data.exported_by} on {evidence_data.export_date.strftime('%Y-%m-%d')}",
            )

            # Title
            self._add_title(doc, "Evidence Inventory Report")

            # Case metadata
            p = doc.add_paragraph()
            p.add_run(f"Case: {evidence_data.case_title}").bold = True
            p.space_after = Pt(5)

            p = doc.add_paragraph()
            p.add_run(f"Total Evidence Items: {evidence_data.total_items}").bold = True
            p.space_after = Pt(10)

            # Evidence list
            self._add_evidence_section(doc, evidence_data.evidence)

            # Log audit event
            if self.audit_logger:
                await self._log_audit(
                    "evidence_list_generated", evidence_data.case_id, evidence_data.exported_by
                )

            return self._document_to_bytes(doc)

        except Exception as exc:
            if self.audit_logger:
                await self._log_audit(
                    "evidence_list_failed",
                    evidence_data.case_id,
                    evidence_data.exported_by,
                    error=str(e),
                )
            raise

    async def generate_timeline_report(self, timeline_data: TimelineExportData) -> bytes:
        """
        Generate timeline report DOCX document.

        Creates chronological timeline of case events including:
        - Case identification
        - Total event count
        - Detailed event entries with dates and descriptions

        Args:
            timeline_data: Timeline export data structure

        Returns:
            bytes: DOCX document as bytes buffer

        Example:
            ```python
            timeline_data = TimelineExportData(
                case_id=1,
                case_title="Smith v. Jones",
                events=[...],
                export_date=datetime.now(),
                exported_by="user@example.com"
            )

            docx_bytes = await generator.generate_timeline_report(timeline_data)
            ```
        """
        try:
            doc = Document()
            self._set_margins(doc)

            # Add header and footer
            self._add_header(doc, f"Timeline Report: {timeline_data.case_title}")
            self._add_footer(
                doc,
                f"Exported by {timeline_data.exported_by} on {timeline_data.export_date.strftime('%Y-%m-%d')}",
            )

            # Title
            self._add_title(doc, "Timeline Report")

            # Case metadata
            p = doc.add_paragraph()
            p.add_run(f"Case: {timeline_data.case_title}").bold = True
            p.space_after = Pt(5)

            p = doc.add_paragraph()
            p.add_run(f"Total Events: {len(timeline_data.events)}").bold = True
            p.space_after = Pt(10)

            # Timeline section
            self._add_timeline_section(doc, timeline_data.events)

            # Log audit event
            if self.audit_logger:
                await self._log_audit(
                    "timeline_report_generated", timeline_data.case_id, timeline_data.exported_by
                )

            return self._document_to_bytes(doc)

        except Exception as exc:
            if self.audit_logger:
                await self._log_audit(
                    "timeline_report_failed",
                    timeline_data.case_id,
                    timeline_data.exported_by,
                    error=str(e),
                )
            raise

    async def generate_case_notes(self, notes_data: NotesExportData) -> bytes:
        """
        Generate case notes report DOCX document.

        Creates organized collection of case notes including:
        - Case identification
        - Total notes count
        - Individual notes with titles, content, and timestamps

        Args:
            notes_data: Notes export data structure

        Returns:
            bytes: DOCX document as bytes buffer

        Example:
            ```python
            notes_data = NotesExportData(
                case_id=1,
                case_title="Smith v. Jones",
                notes=[...],
                export_date=datetime.now(),
                exported_by="user@example.com",
                total_notes=5
            )

            docx_bytes = await generator.generate_case_notes(notes_data)
            ```
        """
        try:
            doc = Document()
            self._set_margins(doc)

            # Add header and footer
            self._add_header(doc, f"Case Notes: {notes_data.case_title}")
            self._add_footer(
                doc,
                f"Exported by {notes_data.exported_by} on {notes_data.export_date.strftime('%Y-%m-%d')}",
            )

            # Title
            self._add_title(doc, "Case Notes Report")

            # Case metadata
            p = doc.add_paragraph()
            p.add_run(f"Case: {notes_data.case_title}").bold = True
            p.space_after = Pt(5)

            p = doc.add_paragraph()
            p.add_run(f"Total Notes: {notes_data.total_notes}").bold = True
            p.space_after = Pt(10)

            # Notes section
            self._add_notes_section(doc, notes_data.notes)

            # Log audit event
            if self.audit_logger:
                await self._log_audit(
                    "case_notes_generated", notes_data.case_id, notes_data.exported_by
                )

            return self._document_to_bytes(doc)

        except Exception as exc:
            if self.audit_logger:
                await self._log_audit(
                    "case_notes_failed", notes_data.case_id, notes_data.exported_by, error=str(e)
                )
            raise

    # Private helper methods

    def _set_margins(self, doc: Document) -> None:
        """Set document margins (1 inch on all sides)."""
        sections = doc.sections
        for section in sections:
            section.top_margin = self.MARGIN_TOP
            section.bottom_margin = self.MARGIN_BOTTOM
            section.left_margin = self.MARGIN_LEFT
            section.right_margin = self.MARGIN_RIGHT

    def _add_header(self, doc: Document, text: str) -> None:
        """
        Add header to document with right-aligned bold text.

        Args:
            doc: Document instance
            text: Header text content
        """
        section = doc.sections[0]
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = text
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.runs[0] if p.runs else p.add_run(text)
        run.bold = True
        run.font.size = self.FONT_SIZE_BODY

    def _add_footer(self, doc: Document, export_info: str) -> None:
        """
        Add footer to document with export info and page numbers (center-aligned).

        Args:
            doc: Document instance
            export_info: Export metadata text
        """
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add export info
        run = p.add_run(export_info)
        run.font.size = self.FONT_SIZE_FOOTER

        # Add separator
        run = p.add_run(" | Page ")
        run.font.size = self.FONT_SIZE_FOOTER

        # Add page number field
        self._add_page_number(p)

        # Add " of " text
        run = p.add_run(" of ")
        run.font.size = self.FONT_SIZE_FOOTER

        # Add total pages field
        self._add_total_pages(p)

    def _add_page_number(self, paragraph) -> None:
        """Add current page number field to paragraph."""
        run = paragraph.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.size = self.FONT_SIZE_FOOTER

    def _add_total_pages(self, paragraph) -> None:
        """Add total pages field to paragraph."""
        run = paragraph.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "NUMPAGES"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.size = self.FONT_SIZE_FOOTER

    def _add_title(self, doc: Document, title: str) -> None:
        """
        Add main title (Heading 1) to document.

        Args:
            doc: Document instance
            title: Title text
        """
        p = doc.add_heading(title, level=1)
        p.space_after = self.SPACING_AFTER_TITLE
        # Heading 1 is automatically bold and larger

    def _add_case_info(self, doc: Document, case_data: CaseExportData) -> None:
        """
        Add case information section.

        Args:
            doc: Document instance
            case_data: Case export data
        """
        # Case ID
        p = doc.add_paragraph()
        p.add_run(f"Case ID: {case_data.case.id}")
        p.space_after = self.SPACING_AFTER_PARAGRAPH

        # Description
        p = doc.add_paragraph()
        p.add_run(f"Description: {case_data.case.description}")
        p.space_after = self.SPACING_AFTER_PARAGRAPH

        # Status
        p = doc.add_paragraph()
        p.add_run(f"Status: {case_data.case.status}")
        p.space_after = Pt(10)

    def _add_evidence_section(self, doc: Document, evidence_list: List[Evidence]) -> None:
        """
        Add evidence inventory section.

        Args:
            doc: Document instance
            evidence_list: List of evidence items
        """
        # Section heading
        p = doc.add_heading("Evidence", level=2)
        p.space_after = self.SPACING_AFTER_HEADING

        for evidence in evidence_list:
            # Evidence title (Heading 3)
            p = doc.add_heading(evidence.title, level=3)
            p.space_after = self.SPACING_AFTER_PARAGRAPH

            # Evidence type
            p = doc.add_paragraph()
            p.add_run(f"Type: {evidence.evidence_type}")
            p.space_after = self.SPACING_AFTER_PARAGRAPH

            # Obtained date (if present)
            if evidence.obtained_date:
                try:
                    obtained_date = datetime.fromisoformat(
                        evidence.obtained_date.replace("Z", "+00:00")
                    )
                    p = doc.add_paragraph()
                    p.add_run(f"Date Obtained: {obtained_date.strftime('%Y-%m-%d')}")
                    p.space_after = Pt(10)
                except ValueError:
                    # Skip invalid dates
                    pass

    def _add_timeline_section(self, doc: Document, timeline_items: List[TimelineEvent]) -> None:
        """
        Add timeline section with chronological events.

        Args:
            doc: Document instance
            timeline_items: List of timeline events
        """
        # Section heading
        p = doc.add_heading("Timeline", level=2)
        p.space_after = self.SPACING_AFTER_HEADING

        for item in timeline_items:
            # Event title (Heading 3)
            p = doc.add_heading(item.title, level=3)
            p.space_after = self.SPACING_AFTER_PARAGRAPH

            # Description (if present)
            if item.description:
                p = doc.add_paragraph()
                p.add_run(item.description)
                p.space_after = self.SPACING_AFTER_PARAGRAPH

            # Event date
            try:
                event_date = datetime.fromisoformat(item.event_date.replace("Z", "+00:00"))
                p = doc.add_paragraph()
                p.add_run(f"Event Date: {event_date.strftime('%Y-%m-%d')}")
                p.space_after = Pt(10)
            except ValueError:
                # Skip invalid dates
                pass

    def _add_notes_section(self, doc: Document, notes: List[Note]) -> None:
        """
        Add notes section with all case notes.

        Args:
            doc: Document instance
            notes: List of notes
        """
        # Section heading
        p = doc.add_heading("Notes", level=2)
        p.space_after = self.SPACING_AFTER_HEADING

        for note in notes:
            # Note title (Heading 3)
            title = note.title or "Untitled Note"
            p = doc.add_heading(title, level=3)
            p.space_after = self.SPACING_AFTER_PARAGRAPH

            # Note content
            p = doc.add_paragraph()
            p.add_run(note.content)
            p.space_after = self.SPACING_AFTER_PARAGRAPH

            # Created date
            try:
                created_date = datetime.fromisoformat(note.created_at.replace("Z", "+00:00"))
                p = doc.add_paragraph()
                p.add_run(f"Created: {created_date.strftime('%Y-%m-%d')}")
                p.space_after = Pt(10)
            except ValueError:
                # Skip invalid dates
                pass

    def _document_to_bytes(self, doc: Document) -> bytes:
        """
        Convert Document object to bytes buffer.

        Args:
            doc: Document instance

        Returns:
            bytes: DOCX file as bytes
        """
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    async def _log_audit(
        self, action: str, case_id: int, user: str, error: Optional[str] = None
    ) -> None:
        """
        Log audit event for export operation.

        Args:
            action: Action type (e.g., 'case_summary_generated')
            case_id: Case identifier
            user: User who performed export
            error: Optional error message if operation failed
        """
        if self.audit_logger and hasattr(self.audit_logger, "log"):
            try:
                await self.audit_logger.log(
                    action=action,
                    resource_type="case",
                    resource_id=case_id,
                    user=user,
                    metadata={"error": error} if error else None,
                )
            except Exception:
                # Don't let audit logging errors break exports
                pass
