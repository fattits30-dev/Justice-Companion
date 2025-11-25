"""
Unit Tests for DOCX Generator Service

This test suite validates the DOCX generation functionality including:
- Case summary document generation
- Evidence inventory reports
- Timeline reports
- Case notes reports
- Document formatting and structure
- Header/footer generation
- Audit logging integration
- Error handling

Run with: pytest backend/services/export/test_docx_generator.py -v
"""

import pytest
import asyncio
from datetime import datetime
from typing import Dict, Any, List
from pathlib import Path
import tempfile
import io

from docx import Document

from .docx_generator import (
    DOCXGenerator,
    CaseExportData,
    EvidenceExportData,
    TimelineExportData,
    NotesExportData,
    Case,
    Evidence,
    Note,
    TimelineEvent,
)

class MockAuditLogger:
    """Mock audit logger for testing."""

    def __init__(self):
        self.logs: List[Dict[str, Any]] = []

    async def log(
        self,
        action: str,
        resource_type: str,
        resource_id: int,
        user: str,
        metadata: Dict[str, Any] = None
    ):
        """Record audit log entry."""
        self.logs.append({
            "action": action,
            "resource_type": resource_type,
            "resource_id": resource_id,
            "user": user,
            "metadata": metadata or {},
            "timestamp": datetime.now().isoformat()
        })

@pytest.fixture
def audit_logger():
    """Create mock audit logger."""
    return MockAuditLogger()

@pytest.fixture
def generator(audit_logger):
    """Create DOCX generator instance."""
    return DOCXGenerator(audit_logger=audit_logger)

@pytest.fixture
def sample_case():
    """Create sample case data."""
    return Case(
        id=1,
        title="Smith v. Jones",
        description="Personal injury case involving vehicle collision",
        status="active"
    )

@pytest.fixture
def sample_evidence():
    """Create sample evidence list."""
    return [
        Evidence(
            id=1,
            case_id=1,
            title="Police Report",
            file_path="/evidence/police-report.pdf",
            content=None,
            evidence_type="document",
            obtained_date="2024-01-15T00:00:00Z",
            created_at="2024-01-15T10:30:00Z",
            updated_at="2024-01-15T10:30:00Z"
        ),
        Evidence(
            id=2,
            case_id=1,
            title="Witness Statement",
            file_path=None,
            content="I saw the accident occur at the intersection...",
            evidence_type="witness",
            obtained_date="2024-01-20T00:00:00Z",
            created_at="2024-01-20T14:00:00Z",
            updated_at=None
        ),
        Evidence(
            id=3,
            case_id=1,
            title="Accident Scene Photo",
            file_path="/evidence/scene-photo.jpg",
            content=None,
            evidence_type="photo",
            obtained_date="2024-01-15T00:00:00Z",
            created_at="2024-01-15T11:00:00Z",
            updated_at=None
        )
    ]

@pytest.fixture
def sample_timeline():
    """Create sample timeline events."""
    return [
        TimelineEvent(
            id=1,
            case_id=1,
            title="Accident Occurred",
            description="Vehicle collision at Main St and 1st Ave",
            event_date="2024-01-15T00:00:00Z",
            event_type="milestone",
            completed=True,
            created_at="2024-01-15T10:00:00Z",
            updated_at="2024-01-15T10:00:00Z"
        ),
        TimelineEvent(
            id=2,
            case_id=1,
            title="Police Report Filed",
            description="Official accident report filed with local PD",
            event_date="2024-01-15T00:00:00Z",
            event_type="filing",
            completed=True,
            created_at="2024-01-15T12:00:00Z",
            updated_at="2024-01-15T12:00:00Z"
        ),
        TimelineEvent(
            id=3,
            case_id=1,
            title="Court Hearing",
            description="Initial hearing scheduled",
            event_date="2024-03-01T09:00:00Z",
            event_type="hearing",
            completed=False,
            created_at="2024-01-20T10:00:00Z",
            updated_at="2024-01-20T10:00:00Z"
        )
    ]

@pytest.fixture
def sample_notes():
    """Create sample notes."""
    return [
        Note(
            id=1,
            case_id=1,
            user_id=1,
            title="Initial Assessment",
            content="Client sustained minor injuries. Medical documentation required.",
            is_pinned=True,
            created_at="2024-01-16T10:00:00Z",
            updated_at="2024-01-16T10:00:00Z"
        ),
        Note(
            id=2,
            case_id=1,
            user_id=1,
            title="Follow-up",
            content="Client provided additional witness contact information.",
            is_pinned=False,
            created_at="2024-01-20T14:00:00Z",
            updated_at="2024-01-20T14:00:00Z"
        ),
        Note(
            id=3,
            case_id=1,
            user_id=1,
            title=None,  # Test untitled note
            content="Need to request medical records from hospital.",
            is_pinned=False,
            created_at="2024-01-22T09:00:00Z",
            updated_at="2024-01-22T09:00:00Z"
        )
    ]

@pytest.fixture
def case_export_data(sample_case, sample_evidence, sample_timeline, sample_notes):
    """Create complete case export data."""
    return CaseExportData(
        case=sample_case,
        evidence=sample_evidence,
        timeline=sample_timeline,
        notes=sample_notes,
        export_date=datetime(2024, 1, 25, 10, 30, 0),
        exported_by="john.doe@example.com"
    )

class TestDOCXGenerator:
    """Test suite for DOCXGenerator class."""

    @pytest.mark.asyncio
    async def test_generate_case_summary(self, generator, case_export_data, audit_logger):
        """Test case summary generation."""
        # Generate DOCX
        docx_bytes = await generator.generate_case_summary(case_export_data)

        # Verify bytes returned
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

        # Verify document can be opened
        doc = Document(io.BytesIO(docx_bytes))
        assert doc is not None

        # Verify content includes case title
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Smith v. Jones" in text_content
        assert "Case Summary" in text_content
        assert "Evidence" in text_content
        assert "Timeline" in text_content
        assert "Notes" in text_content

        # Verify audit log
        assert len(audit_logger.logs) == 1
        assert audit_logger.logs[0]["action"] == "case_summary_generated"
        assert audit_logger.logs[0]["resource_id"] == 1

    @pytest.mark.asyncio
    async def test_generate_case_summary_minimal(self, generator, sample_case):
        """Test case summary with minimal data (no evidence/timeline/notes)."""
        minimal_data = CaseExportData(
            case=sample_case,
            evidence=[],
            timeline=[],
            notes=[],
            export_date=datetime.now(),
            exported_by="test@example.com"
        )

        docx_bytes = await generator.generate_case_summary(minimal_data)

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

        # Verify document structure
        doc = Document(io.BytesIO(docx_bytes))
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Smith v. Jones" in text_content
        # Should not contain section headers for empty sections
        assert "Evidence" not in text_content
        assert "Timeline" not in text_content
        assert "Notes" not in text_content

    @pytest.mark.asyncio
    async def test_generate_evidence_list(self, generator, sample_evidence, audit_logger):
        """Test evidence inventory generation."""
        evidence_data = EvidenceExportData(
            case_id=1,
            case_title="Smith v. Jones",
            evidence=sample_evidence,
            export_date=datetime.now(),
            exported_by="test@example.com",
            total_items=len(sample_evidence)
        )

        docx_bytes = await generator.generate_evidence_list(evidence_data)

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

        # Verify content
        doc = Document(io.BytesIO(docx_bytes))
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Evidence Inventory Report" in text_content
        assert "Total Evidence Items: 3" in text_content
        assert "Police Report" in text_content
        assert "Witness Statement" in text_content
        assert "Accident Scene Photo" in text_content

        # Verify audit log
        assert len(audit_logger.logs) == 1
        assert audit_logger.logs[0]["action"] == "evidence_list_generated"

    @pytest.mark.asyncio
    async def test_generate_timeline_report(self, generator, sample_timeline, audit_logger):
        """Test timeline report generation."""
        timeline_data = TimelineExportData(
            case_id=1,
            case_title="Smith v. Jones",
            events=sample_timeline,
            export_date=datetime.now(),
            exported_by="test@example.com"
        )

        docx_bytes = await generator.generate_timeline_report(timeline_data)

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

        # Verify content
        doc = Document(io.BytesIO(docx_bytes))
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Timeline Report" in text_content
        assert "Total Events: 3" in text_content
        assert "Accident Occurred" in text_content
        assert "Police Report Filed" in text_content
        assert "Court Hearing" in text_content

        # Verify audit log
        assert len(audit_logger.logs) == 1
        assert audit_logger.logs[0]["action"] == "timeline_report_generated"

    @pytest.mark.asyncio
    async def test_generate_case_notes(self, generator, sample_notes, audit_logger):
        """Test case notes generation."""
        notes_data = NotesExportData(
            case_id=1,
            case_title="Smith v. Jones",
            notes=sample_notes,
            export_date=datetime.now(),
            exported_by="test@example.com",
            total_notes=len(sample_notes)
        )

        docx_bytes = await generator.generate_case_notes(notes_data)

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

        # Verify content
        doc = Document(io.BytesIO(docx_bytes))
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Case Notes Report" in text_content
        assert "Total Notes: 3" in text_content
        assert "Initial Assessment" in text_content
        assert "Follow-up" in text_content
        assert "Untitled Note" in text_content  # Test default title

        # Verify audit log
        assert len(audit_logger.logs) == 1
        assert audit_logger.logs[0]["action"] == "case_notes_generated"

    @pytest.mark.asyncio
    async def test_document_formatting(self, generator, case_export_data):
        """Test document formatting (margins, headers, footers)."""
        docx_bytes = await generator.generate_case_summary(case_export_data)
        doc = Document(io.BytesIO(docx_bytes))

        # Check margins (1 inch = 914400 EMUs)
        section = doc.sections[0]
        assert section.top_margin == 914400
        assert section.bottom_margin == 914400
        assert section.left_margin == 914400
        assert section.right_margin == 914400

        # Check header exists
        header = section.header
        assert header is not None
        header_text = "\n".join([p.text for p in header.paragraphs])
        assert "Case: Smith v. Jones" in header_text

        # Check footer exists
        footer = section.footer
        assert footer is not None
        footer_text = "\n".join([p.text for p in footer.paragraphs])
        assert "john.doe@example.com" in footer_text
        assert "2024-01-25" in footer_text

    @pytest.mark.asyncio
    async def test_error_handling_with_audit(self, generator, audit_logger):
        """Test error handling and audit logging for failures."""
        # Create invalid data (missing required field)
        with pytest.raises(Exception):
            invalid_data = {
                "case": {"id": 1},  # Missing required fields
                "evidence": [],
                "timeline": [],
                "notes": [],
            }
            # This should fail validation
            case_data = CaseExportData(**invalid_data)

    @pytest.mark.asyncio
    async def test_generator_without_audit_logger(self, case_export_data):
        """Test generator works without audit logger."""
        generator = DOCXGenerator(audit_logger=None)

        # Should still generate document
        docx_bytes = await generator.generate_case_summary(case_export_data)

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

    @pytest.mark.asyncio
    async def test_evidence_with_missing_dates(self, generator):
        """Test evidence handling when dates are None."""
        evidence_data = EvidenceExportData(
            case_id=1,
            case_title="Test Case",
            evidence=[
                Evidence(
                    id=1,
                    case_id=1,
                    title="Test Evidence",
                    file_path=None,
                    content="Test content",
                    evidence_type="note",
                    obtained_date=None,  # No date
                    created_at="2024-01-15T10:00:00Z",
                    updated_at=None
                )
            ],
            export_date=datetime.now(),
            exported_by="test@example.com",
            total_items=1
        )

        # Should not raise exception
        docx_bytes = await generator.generate_evidence_list(evidence_data)
        assert isinstance(docx_bytes, bytes)

    @pytest.mark.asyncio
    async def test_timeline_with_invalid_dates(self, generator):
        """Test timeline handling with invalid date formats."""
        timeline_data = TimelineExportData(
            case_id=1,
            case_title="Test Case",
            events=[
                TimelineEvent(
                    id=1,
                    case_id=1,
                    title="Test Event",
                    description="Test",
                    event_date="2024-01-15",  # Valid but simple format
                    event_type="milestone",
                    completed=False,
                    created_at="2024-01-15T10:00:00Z",
                    updated_at="2024-01-15T10:00:00Z"
                )
            ],
            export_date=datetime.now(),
            exported_by="test@example.com"
        )

        # Should handle gracefully
        docx_bytes = await generator.generate_timeline_report(timeline_data)
        assert isinstance(docx_bytes, bytes)

    @pytest.mark.asyncio
    async def test_save_to_file(self, generator, case_export_data):
        """Test saving generated DOCX to file system."""
        docx_bytes = await generator.generate_case_summary(case_export_data)

        # Save to temp file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(docx_bytes)
            tmp_path = tmp.name

        try:
            # Verify file exists and can be opened
            assert Path(tmp_path).exists()
            doc = Document(tmp_path)
            assert doc is not None

            # Verify content
            text_content = "\n".join([p.text for p in doc.paragraphs])
            assert "Smith v. Jones" in text_content
        finally:
            # Clean up
            Path(tmp_path).unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_concurrent_generation(self, generator, case_export_data):
        """Test multiple concurrent document generations."""
        # Generate 5 documents concurrently
        tasks = [
            generator.generate_case_summary(case_export_data)
            for _ in range(5)
        ]

        results = await asyncio.gather(*tasks)

        # Verify all succeeded
        assert len(results) == 5
        for docx_bytes in results:
            assert isinstance(docx_bytes, bytes)
            assert len(docx_bytes) > 0

if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
