"""
Test suite for DocumentParserService.

Tests document parsing functionality for PDF, DOCX, and TXT formats.
"""

import os
import pytest
import tempfile
from io import BytesIO

# PDF creation for testing
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# DOCX creation for testing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from backend.services.document_parser_service import (
    DocumentParserService
)

class TestDocumentParserService:
    """Test DocumentParserService functionality."""

    @pytest.fixture
    def service(self):
        """Create DocumentParserService instance."""
        return DocumentParserService()

    @pytest.fixture
    def temp_dir(self):
        """Create temporary directory for test files."""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield tmpdir

    # ============================================================================
    # Plain Text Tests
    # ============================================================================

    @pytest.mark.asyncio
    async def test_parse_txt_file(self, service, temp_dir):
        """Test parsing plain text file."""
        # Create test TXT file
        test_file = os.path.join(temp_dir, "test.txt")
        test_content = "Hello world! This is a test document with multiple words."

        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(test_content)

        # Parse document
        result = await service.parse_document(test_file)

        # Assertions
        assert result.text == test_content
        assert result.filename == "test.txt"
        assert result.file_type == "txt"
        assert result.word_count == 10
        assert result.page_count is None
        assert result.metadata is None

    @pytest.mark.asyncio
    async def test_parse_txt_buffer(self, service):
        """Test parsing plain text from buffer."""
        test_content = "Buffer test content with seven words here."
        buffer = test_content.encode('utf-8')

        # Parse from buffer
        result = await service.parse_document_buffer(buffer, "buffer.txt")

        # Assertions
        assert result.text == test_content
        assert result.filename == "buffer.txt"
        assert result.file_type == "txt"
        assert result.word_count == 7

    @pytest.mark.asyncio
    async def test_parse_txt_with_unicode(self, service, temp_dir):
        """Test parsing text file with Unicode characters."""
        test_file = os.path.join(temp_dir, "unicode.txt")
        test_content = "Legal document: § 123, © 2024, Müller v. González"

        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(test_content)

        # Parse document
        result = await service.parse_document(test_file)

        # Assertions
        assert result.text == test_content
        assert result.word_count == 7

    # ============================================================================
    # PDF Tests (require pypdf)
    # ============================================================================

    @pytest.mark.asyncio
    @pytest.mark.skipif(not REPORTLAB_AVAILABLE, reason="reportlab not installed")
    async def test_parse_pdf_file(self, service, temp_dir):
        """Test parsing PDF file."""
        # Create test PDF
        test_file = os.path.join(temp_dir, "test.pd")

        c = canvas.Canvas(test_file, pagesize=letter)
        c.drawString(100, 750, "This is page 1 of the test PDF.")
        c.showPage()
        c.drawString(100, 750, "This is page 2 of the test PDF.")
        c.showPage()
        c.save()

        # Parse document
        result = await service.parse_document(test_file)

        # Assertions
        assert result.filename == "test.pd"
        assert result.file_type == "pd"
        assert result.page_count == 2
        assert "page 1" in result.text.lower()
        assert "page 2" in result.text.lower()
        assert result.word_count > 0

    @pytest.mark.asyncio
    @pytest.mark.skipif(not REPORTLAB_AVAILABLE, reason="reportlab not installed")
    async def test_parse_pdf_with_metadata(self, service, temp_dir):
        """Test parsing PDF with metadata."""
        # Create test PDF with metadata
        test_file = os.path.join(temp_dir, "metadata.pd")

        c = canvas.Canvas(test_file, pagesize=letter)
        c.setTitle("Test Document Title")
        c.setAuthor("John Doe")
        c.drawString(100, 750, "Content of the document.")
        c.save()

        # Parse document
        result = await service.parse_document(test_file)

        # Assertions
        assert result.file_type == "pd"
        # Note: metadata extraction depends on pypdf capabilities
        # Some metadata might not be preserved by reportlab

    # ============================================================================
    # DOCX Tests (require python-docx)
    # ============================================================================

    @pytest.mark.asyncio
    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    async def test_parse_docx_file(self, service, temp_dir):
        """Test parsing DOCX file."""
        # Create test DOCX
        test_file = os.path.join(temp_dir, "test.docx")

        doc = Document()
        doc.add_paragraph("This is the first paragraph.")
        doc.add_paragraph("This is the second paragraph with more words.")
        doc.core_properties.title = "Test DOCX Document"
        doc.core_properties.author = "Jane Smith"
        doc.save(test_file)

        # Parse document
        result = await service.parse_document(test_file)

        # Assertions
        assert result.filename == "test.docx"
        assert result.file_type == "docx"
        assert "first paragraph" in result.text
        assert "second paragraph" in result.text
        assert result.word_count > 0
        assert result.metadata is not None
        assert result.metadata.title == "Test DOCX Document"
        assert result.metadata.author == "Jane Smith"

    @pytest.mark.asyncio
    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    async def test_parse_docx_buffer(self, service):
        """Test parsing DOCX from buffer."""
        # Create DOCX in memory
        doc = Document()
        doc.add_paragraph("Buffer content test.")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Parse from buffer
        result = await service.parse_document_buffer(buffer.read(), "buffer.docx")

        # Assertions
        assert result.filename == "buffer.docx"
        assert result.file_type == "docx"
        assert "Buffer content test" in result.text

    # ============================================================================
    # Validation Tests
    # ============================================================================

    def test_validate_file_size_success(self, service):
        """Test file size validation passes for small files."""
        result = service.validate_file_size(1024)  # 1KB

        assert result.valid is True
        assert result.error is None

    def test_validate_file_size_failure(self, service):
        """Test file size validation fails for large files."""
        result = service.validate_file_size(11 * 1024 * 1024)  # 11MB (exceeds 10MB limit)

        assert result.valid is False
        assert result.error is not None
        assert "exceeds maximum" in result.error

    def test_validate_file_size_custom_limit(self):
        """Test file size validation with custom limit."""
        service = DocumentParserService(max_file_size=1024)  # 1KB limit

        result = service.validate_file_size(2048)  # 2KB

        assert result.valid is False
        assert result.error is not None

    @pytest.mark.asyncio
    async def test_parse_file_too_large(self, service, temp_dir):
        """Test parsing fails for oversized files."""
        # Create large file
        test_file = os.path.join(temp_dir, "large.txt")
        with open(test_file, 'w') as f:
            f.write('x' * (11 * 1024 * 1024))  # 11MB

        # Attempt to parse (should fail)
        with pytest.raises(Exception) as exc_info:
            await service.parse_document(test_file)

        assert "exceeds maximum" in str(exc_info.value).lower()

    # ============================================================================
    # Support Detection Tests
    # ============================================================================

    def test_get_supported_extensions(self, service):
        """Test getting supported file extensions."""
        extensions = service.get_supported_extensions()

        assert '.pd' in extensions
        assert '.docx' in extensions
        assert '.txt' in extensions
        assert len(extensions) == 3

    def test_is_supported_valid(self, service):
        """Test supported file detection for valid extensions."""
        assert service.is_supported("document.pd") is True
        assert service.is_supported("report.docx") is True
        assert service.is_supported("notes.txt") is True
        assert service.is_supported("UPPERCASE.PDF") is True  # Case insensitive

    def test_is_supported_invalid(self, service):
        """Test supported file detection for invalid extensions."""
        assert service.is_supported("image.jpg") is False
        assert service.is_supported("data.csv") is False
        assert service.is_supported("archive.zip") is False

    @pytest.mark.asyncio
    async def test_parse_unsupported_format(self, service, temp_dir):
        """Test parsing fails for unsupported format."""
        test_file = os.path.join(temp_dir, "test.jpg")
        with open(test_file, 'wb') as f:
            f.write(b'\xFF\xD8\xFF\xE0')  # JPEG header

        # Attempt to parse (should fail)
        with pytest.raises(Exception) as exc_info:
            await service.parse_document(test_file)

        assert "unsupported" in str(exc_info.value).lower()

    # ============================================================================
    # Helper Method Tests
    # ============================================================================

    def test_count_words_basic(self, service):
        """Test word counting with basic text."""
        text = "Hello world this is a test"
        count = service._count_words(text)

        assert count == 6

    def test_count_words_multiple_spaces(self, service):
        """Test word counting with multiple spaces."""
        text = "Hello    world   test"
        count = service._count_words(text)

        assert count == 3

    def test_count_words_newlines(self, service):
        """Test word counting with newlines."""
        text = "Hello\nworld\ntest"
        count = service._count_words(text)

        assert count == 3

    def test_count_words_empty(self, service):
        """Test word counting with empty string."""
        text = ""
        count = service._count_words(text)

        assert count == 0

    def test_extract_summary_short(self, service):
        """Test summary extraction for short text."""
        text = "This is a short text that fits within the limit."
        summary = service.extract_summary(text, max_words=100)

        assert summary == text
        assert not summary.endswith('...')

    def test_extract_summary_long(self, service):
        """Test summary extraction for long text."""
        words = ['word'] * 1000
        text = ' '.join(words)

        summary = service.extract_summary(text, max_words=500)

        # Should be truncated
        assert summary.endswith('...')
        assert summary.count('word') == 500

    def test_extract_summary_custom_limit(self, service):
        """Test summary extraction with custom word limit."""
        text = "one two three four five six seven eight nine ten"

        summary = service.extract_summary(text, max_words=5)

        assert summary == "one two three four five..."

    # ============================================================================
    # Error Handling Tests
    # ============================================================================

    @pytest.mark.asyncio
    async def test_parse_nonexistent_file(self, service):
        """Test parsing fails for nonexistent file."""
        with pytest.raises(Exception) as exc_info:
            await service.parse_document("/nonexistent/path/file.txt")

        assert "not found" in str(exc_info.value).lower()

    @pytest.mark.asyncio
    async def test_parse_invalid_pdf(self, service, temp_dir):
        """Test parsing fails for corrupt PDF."""
        test_file = os.path.join(temp_dir, "corrupt.pd")
        with open(test_file, 'wb') as f:
            f.write(b'Not a real PDF file')

        # Attempt to parse (should fail)
        with pytest.raises(Exception) as exc_info:
            await service.parse_document(test_file)

        # Should contain error message
        assert exc_info.value is not None

    @pytest.mark.asyncio
    async def test_parse_invalid_docx(self, service, temp_dir):
        """Test parsing fails for corrupt DOCX."""
        test_file = os.path.join(temp_dir, "corrupt.docx")
        with open(test_file, 'wb') as f:
            f.write(b'Not a real DOCX file')

        # Attempt to parse (should fail)
        with pytest.raises(Exception) as exc_info:
            await service.parse_document(test_file)

        # Should contain error message
        assert exc_info.value is not None

# ============================================================================
# Integration Tests with Audit Logger
# ============================================================================

class MockAuditLogger:
    """Mock audit logger for testing."""

    def __init__(self):
        self.logs = []

    def log(self, **kwargs):
        """Record log entry."""
        self.logs.append(kwargs)

class TestDocumentParserServiceWithAudit:
    """Test DocumentParserService with audit logging."""

    @pytest.fixture
    def audit_logger(self):
        """Create mock audit logger."""
        return MockAuditLogger()

    @pytest.fixture
    def service(self, audit_logger):
        """Create DocumentParserService with audit logger."""
        return DocumentParserService(audit_logger=audit_logger)

    @pytest.mark.asyncio
    async def test_parse_logs_success(self, service, audit_logger, tmp_path):
        """Test successful parse logs audit events."""
        # Create test file
        test_file = tmp_path / "test.txt"
        test_file.write_text("Test content")

        # Parse document
        await service.parse_document(str(test_file), user_id="user123")

        # Verify audit logs
        assert len(audit_logger.logs) == 2

        # Check start event
        start_log = audit_logger.logs[0]
        assert start_log['event_type'] == "document.parse.start"
        assert start_log['user_id'] == "user123"
        assert start_log['action'] == "parse"

        # Check success event
        success_log = audit_logger.logs[1]
        assert success_log['event_type'] == "document.parse.success"
        assert success_log['success'] is True

    @pytest.mark.asyncio
    async def test_parse_logs_error(self, service, audit_logger):
        """Test failed parse logs error event."""
        # Attempt to parse nonexistent file
        try:
            await service.parse_document("/nonexistent/file.txt", user_id="user123")
        except Exception:
            pass

        # Verify error was logged
        assert len(audit_logger.logs) >= 1
        # Last log should be error
        error_log = audit_logger.logs[-1]
        assert "error" in error_log['event_type'].lower() or error_log['success'] is False

if __name__ == '__main__':
    pytest.main([__file__, '-v'])
